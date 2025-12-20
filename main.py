from fastapi import FastAPI, UploadFile, Form
from fastapi.responses import FileResponse, HTMLResponse
import fitz
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import os
import uuid
import argostranslate.package
import argostranslate.translate
import urllib.request

# Configure OCR path
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\poppler\Library\bin"

app = FastAPI()

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# -------------------------------
# Automatic EN ↔ KM model installer
# -------------------------------
def ensure_en_km_model():
    installed_languages = argostranslate.translate.get_installed_languages()
    codes = [lang.code for lang in installed_languages]
    
    if "en" in codes and "km" in codes:
        # Check if translation exists
        for lang in installed_languages:
            for to_lang in installed_languages:
                if lang.code == "en" and to_lang.code == "km":
                    return  # Model exists
        # Otherwise, fall through to download
    print("[INFO] Downloading EN↔KM Argos model...")
    # URL to the Argos Translate EN→KM model
    url = "https://argosopentech.com/models/translate-en_km.argosmodel"
    local_path = "translate-en_km.argosmodel"
    
    # Download if not exists
    if not os.path.exists(local_path):
        urllib.request.urlretrieve(url, local_path)
    
    argostranslate.package.install_from_path(local_path)
    print("[INFO] EN↔KM model installed successfully.")

# Call on startup
ensure_en_km_model()

# Build lang_pairs after ensuring the model
installed_languages = argostranslate.translate.get_installed_languages()
lang_pairs = {}
for from_lang in installed_languages:
    for to_lang in installed_languages:
        if from_lang != to_lang:
            lang_pairs[f"{from_lang.code}-{to_lang.code}"] = from_lang.get_translation(to_lang)

# -------------------------------
# PDF text extraction
# -------------------------------
def extract_text(pdf_path, ocr_langs="eng+khm"):
    doc = fitz.open(pdf_path)
    all_text = ""
    images_cache = None

    for page_num, page in enumerate(doc, start=1):
        page_text = page.get_text().strip()
        if page_text:
            all_text += page_text + "\n"
        else:
            if images_cache is None:
                try:
                    images_cache = convert_from_path(
                        pdf_path, dpi=150, poppler_path=POPPLER_PATH
                    )
                except Exception as e:
                    print(f"[ERROR] Failed to convert PDF to images: {e}")
                    continue
            img = images_cache[page_num - 1]
            try:
                ocr_text = pytesseract.image_to_string(img, lang=ocr_langs)
                all_text += ocr_text + "\n"
            except pytesseract.TesseractError as e:
                print(f"[WARNING] Tesseract OCR failed on page {page_num}: {e}")
            except Exception as e:
                print(f"[WARNING] Unexpected error on page {page_num}: {e}")

    return all_text.strip()

# -------------------------------
# Text translation
# -------------------------------
def translate_text(text, direction):
    translation = lang_pairs.get(direction)
    if not translation:
        return f"[ERROR] No translation model installed for {direction}. Please install it."
    return translation.translate(text)

# -------------------------------
# FastAPI routes
# -------------------------------
@app.get("/", response_class=HTMLResponse)
def index():
    return open("frontend.html", encoding="utf-8").read()

@app.post("/translate")
async def translate_pdf(file: UploadFile, direction: str = Form(...)):
    uid = str(uuid.uuid4())
    pdf_path = f"{UPLOAD_DIR}/{uid}.pdf"
    with open(pdf_path, "wb") as f:
        f.write(await file.read())

    raw_text = extract_text(pdf_path)
    translated = translate_text(raw_text, direction)

    doc = Document()
    doc.add_heading("Translated Document", level=1)
    for line in translated.split("\n"):
        doc.add_paragraph(line)

    output_path = f"{OUTPUT_DIR}/{uid}.docx"
    doc.save(output_path)
    return FileResponse(output_path, filename="translated.docx")
