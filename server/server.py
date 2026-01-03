from fastapi import FastAPI, UploadFile, Form, File, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware

import fitz  # PyMuPDF
from googletrans import Translator, LANGUAGES
from docx import Document
from docx.shared import Pt
import pytesseract
from PIL import Image

import io
import time
import re
import os

app = FastAPI()

# CORS - Allow your frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:5500",
        "http://localhost:5173",
        "https://*.vercel.app",
        "https://pdf-translator-five.vercel.app",
        "https://pdf-translator-mei-romneys-projects.vercel.app",
        "https://pdf-translator-khm-en.vercel.app",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def root():
    return {
        "message": "PDF Translator API - Multi-Language Support",
        "status": "active",
        "endpoints": {
            "translate": "/translate (POST)",
            "languages": "/languages (GET)"
        }
    }

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

@app.get("/languages")
async def get_supported_languages():
    """Return all supported languages"""
    return {
        "languages": LANGUAGES,
        "total": len(LANGUAGES)
    }

def clean_text_for_xml(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    return text.encode("utf-8", errors="ignore").decode("utf-8", errors="ignore")


def extract_text_from_pdf_ocr(pdf_bytes: bytes, lang: str = 'eng') -> str:
    """Extract text from PDF using OCR"""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = ""
    
    # Map common language codes to Tesseract language codes
    tesseract_lang_map = {
        'en': 'eng',
        'km': 'khm',
        'zh-cn': 'chi_sim',
        'zh-tw': 'chi_tra',
        'ja': 'jpn',
        'ko': 'kor',
        'th': 'tha',
        'vi': 'vie',
        'ar': 'ara',
        'hi': 'hin',
        'fr': 'fra',
        'de': 'deu',
        'es': 'spa',
        'it': 'ita',
        'pt': 'por',
        'ru': 'rus',
    }
    
    tesseract_lang = tesseract_lang_map.get(lang, 'eng')
    
    for page_num, page in enumerate(doc):
        print(f"OCR processing page {page_num + 1}/{len(doc)} with language: {tesseract_lang}...")
        
        try:
            # Convert page to image (300 DPI for better OCR)
            mat = fitz.Matrix(300/72, 300/72)
            pix = page.get_pixmap(matrix=mat)
            
            # Convert to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Perform OCR
            page_text = pytesseract.image_to_string(img, lang=tesseract_lang)
            page_text = clean_text_for_xml(page_text)
            text += page_text + "\n\n"
        except Exception as e:
            print(f"OCR error on page {page_num + 1}: {e}")
            # Fall back to standard extraction if OCR fails
            page_text = page.get_text()
            page_text = clean_text_for_xml(page_text)
            text += page_text + "\n\n"
    
    doc.close()
    return text.strip()


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes in memory (standard method)"""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = ""
    for page in doc:
        page_text = page.get_text()
        page_text = clean_text_for_xml(page_text)
        text += page_text + "\n\n"
    doc.close()
    return text.strip()


def detect_language(text: str) -> str:
    """Detect the language of the text"""
    try:
        translator = Translator()
        # Take a sample of text for detection (first 1000 chars)
        sample = text[:1000] if len(text) > 1000 else text
        detection = translator.detect(sample)
        return detection.lang
    except Exception as e:
        print(f"Language detection error: {e}")
        return "en"  # Default to English


def translate_text(text: str, source_lang: str, target_lang: str) -> str:
    """Translate text using googletrans"""
    translator = Translator()
    
    # Split into smaller chunks (googletrans has 15k char limit per request)
    max_chunk = 4500
    paragraphs = text.split("\n")
    chunks = []
    current = ""

    for p in paragraphs:
        if len(current) + len(p) > max_chunk and current:
            chunks.append(current)
            current = p + "\n"
        else:
            current += p + "\n"

    if current:
        chunks.append(current)

    translated_chunks = []

    for i, chunk in enumerate(chunks):
        chunk = clean_text_for_xml(chunk)
        if not chunk.strip():
            translated_chunks.append(chunk)
            continue
            
        # Retry logic
        max_retries = 3
        for attempt in range(max_retries):
            try:
                print(f"Translating chunk {i + 1}/{len(chunks)} (attempt {attempt + 1})...")
                result = translator.translate(chunk, src=source_lang, dest=target_lang)
                translated_chunks.append(clean_text_for_xml(result.text))
                print(f"✓ Chunk {i + 1}/{len(chunks)} translated successfully")
                time.sleep(0.5)  # Rate limiting
                break
            except Exception as e:
                if attempt < max_retries - 1:
                    print(f"Retry {attempt + 1}/{max_retries} for chunk {i + 1}: {str(e)}")
                    time.sleep(2)
                else:
                    print(f"Failed to translate chunk {i + 1}: {str(e)}")
                    raise HTTPException(
                        status_code=500, 
                        detail=f"Translation failed on chunk {i + 1}: {str(e)}"
                    )

    return "\n".join(translated_chunks)


def create_docx(text: str, title: str = "Document") -> io.BytesIO:
    """Create DOCX document in memory and return BytesIO"""
    doc = Document()
    doc.add_heading(title, level=1)

    text = clean_text_for_xml(text)

    for line in text.split("\n"):
        line = line.strip()
        if line:
            try:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(12)
            except Exception as e:
                print(f"DOCX error: {e}")

    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return docx_io

def create_doc(text: str, title: str = "Document") -> io.BytesIO:
    """Create DOC document (Word 97-2003 format) - using DOCX library"""
    doc = Document()
    doc.add_heading(title, level=1)

    text = clean_text_for_xml(text)

    for line in text.split("\n"):
        line = line.strip()
        if line:
            try:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(12)
            except Exception as e:
                print(f"Doc error: {e}")

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_txt(text: str, title: str = "DOCUMENT") -> io.BytesIO:
    """Create plain text file in memory and return BytesIO"""
    text = clean_text_for_xml(text)

    content = f"{title}\n"
    content += "=" * 50 + "\n\n"
    content += text

    txt_io = io.BytesIO()
    txt_io.write(content.encode('utf-8'))
    txt_io.seek(0)
    return txt_io


def get_language_name(lang_code: str) -> str:
    """Get full language name from code"""
    return LANGUAGES.get(lang_code, lang_code.upper())
    

@app.post("/translate")
async def translate_pdf(
    file: UploadFile = File(...),
    source_lang: str = Form(None),  # Optional - will auto-detect if not provided
    target_lang: str = Form(...),
    use_ocr: bool = Form(False),
    format: str = Form("docx")
):
    """
    Translate or extract text from PDF
    
    Parameters:
    - file: PDF file
    - source_lang: Source language code (e.g., 'en', 'km', 'fr'). If None, will auto-detect
    - target_lang: Target language code. Use same as source_lang for extraction only
    - use_ocr: Whether to use OCR for text extraction
    - format: Output format (docx, doc, txt)
    """
    try:
        if not file.filename.endswith('.pdf'):
            raise HTTPException(status_code=400, detail="Only PDF files are allowed")
        
        pdf_bytes = await file.read()
        
        print(f"DEBUG: source_lang={source_lang}, target_lang={target_lang}, use_ocr={use_ocr}, format={format}")
        
        # Extract text
        if use_ocr or source_lang in ['km', 'ar', 'hi', 'th', 'ja', 'ko', 'zh-cn', 'zh-tw']:
            # Use OCR for languages that often have font issues
            ocr_lang = source_lang if source_lang else 'eng'
            print(f"Extracting text using OCR ({ocr_lang})...")
            text = extract_text_from_pdf_ocr(pdf_bytes, lang=ocr_lang)
        else:
            # Use standard extraction
            print("Extracting text (standard method)...")
            text = extract_text_from_pdf(pdf_bytes)
        
        if not text or len(text) < 10:
            raise HTTPException(status_code=400, detail="Failed to extract text from PDF")
        
        # Auto-detect source language if not provided
        if not source_lang or source_lang == 'auto':
            print("Auto-detecting language...")
            source_lang = detect_language(text)
            print(f"Detected language: {source_lang} ({get_language_name(source_lang)})")
        
        # Check if this is extraction-only (same source and target)
        is_extract_only = (source_lang == target_lang)
        
        if is_extract_only:
            # Extract text only, no translation
            final_text = text
            lang_name = get_language_name(source_lang)
            document_title = f"Extracted Text ({lang_name})"
            print(f"Extraction complete! Extracted {len(text)} characters.")
        else:
            # Translation request
            src_name = get_language_name(source_lang)
            tgt_name = get_language_name(target_lang)
            
            print(f"Translating {source_lang} ({src_name}) → {target_lang} ({tgt_name})")
            final_text = translate_text(text, source_lang, target_lang)
            document_title = f"Translated Document ({src_name} → {tgt_name})"
            print("Translation complete!")

        # Create output file
        print(f"Creating {format.upper()} file...")
        if format == "docx":
            file_io = create_docx(final_text, document_title)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = "result.docx"
        elif format == "doc":
            file_io = create_doc(final_text, document_title)
            media_type = "application/msword"
            filename = "result.doc"
        else:  # txt
            file_io = create_txt(final_text, document_title.upper())
            media_type = "text/plain"
            filename = "result.txt"

        print("Process complete! Sending file to client.")

        return StreamingResponse(
            file_io,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        print("Error:", e)
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("server:app", host="0.0.0.0", port=port, reload=False)